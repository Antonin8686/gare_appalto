from __future__ import annotations

import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from .constants import (
    MATRICE_REQUISITI,
    OFFERTA_ECONOMICA,
    RELAZIONE_TECNICA,
    REPORT_PARTECIPABILITA,
    SCHEDA_GARA,
)
from .collector import TenderExportContext


def _setup_document(title: str, subtitle: str = "") -> Document:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    heading = document.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.color.rgb = RGBColor(15, 23, 42)
        run.font.size = Pt(18)

    if subtitle:
        sub = document.add_paragraph(subtitle)
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in sub.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(100, 116, 139)

    document.add_paragraph("")
    return document


def _add_meta_footer(document: Document, exported_at: str) -> None:
    document.add_paragraph("")
    footer = document.add_paragraph(f"Documento generato il {exported_at}")
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(148, 163, 184)


def _add_key_value_table(document: Document, rows: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for index, (label, value) in enumerate(rows):
        table.rows[index].cells[0].text = label
        table.rows[index].cells[1].text = value or "—"
        for paragraph in table.rows[index].cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True


def _save_document(document: Document) -> bytes:
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def export_scheda_docx(ctx: TenderExportContext) -> bytes:
    scheda = ctx.scheda
    document = _setup_document(
        "Scheda gara d'appalto",
        f"CIG {scheda['cig']} · {scheda.get('organization', '')}",
    )
    rows = [
        ("CIG", scheda["cig"]),
        ("CPV", scheda["cpv"]),
        ("Oggetto", scheda.get("oggetto", "")),
        ("Importo", f"€ {scheda['importo']}"),
        ("Scadenza", scheda["scadenza"]),
        ("Stato", scheda["stato"]),
        ("Fase", scheda["fase"]),
        ("Fonte", scheda["source"]),
        ("Priorità", scheda["priorita"]),
        ("Punteggio priorità", str(scheda.get("priority_score", ""))),
        ("Estrazione AI", "Sì" if scheda.get("ai_extracted") else "No"),
    ]
    _add_key_value_table(document, rows)

    formal_rules = scheda.get("formal_rules") or {}
    if formal_rules:
        document.add_heading("Regole formali", level=2)
        for group, items in formal_rules.items():
            if not items:
                continue
            document.add_heading(group.replace("_", " ").title(), level=3)
            for item in items:
                label = item.get("label", "")
                detail = item.get("detail", "")
                document.add_paragraph(f"{label}: {detail}" if label else detail, style="List Bullet")

    _add_meta_footer(document, ctx.exported_at)
    return _save_document(document)


def export_matrix_docx(ctx: TenderExportContext) -> bytes:
    matrix = ctx.matrix
    document = _setup_document(
        "Matrice requisiti aziende",
        f"Gara CIG {matrix.tender_cig}",
    )

    if not matrix.companies:
        document.add_paragraph("Nessuna azienda configurata per la matrice.")
        return _save_document(document)

    headers = ["Requisito", "Categoria", "Tipo", "Soglia"] + [
        company["name"] for company in matrix.companies
    ]
    table = document.add_table(
        rows=1 + len(matrix.requirements),
        cols=len(headers),
    )
    table.style = "Table Grid"
    for col_index, header in enumerate(headers):
        cell = table.rows[0].cells[col_index]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for row_index, requirement in enumerate(matrix.requirements, start=1):
        cells = table.rows[row_index].cells
        cells[0].text = requirement.descrizione[:500]
        cells[1].text = requirement.categoria_label
        cells[2].text = requirement.tipo_label
        cells[3].text = requirement.soglia_minima or "—"
        for col_index, cell_data in enumerate(requirement.cells, start=4):
            cells[col_index].text = f"{cell_data.esito_label}\n{cell_data.motivazione}"

    document.add_paragraph("")
    summary = matrix.summary
    document.add_paragraph(
        f"Riepilogo: {summary.get('soddisfatto', 0)} soddisfatti · "
        f"{summary.get('parzialmente', 0)} parziali · "
        f"{summary.get('non_soddisfatto', 0)} non soddisfatti"
    )
    _add_meta_footer(document, ctx.exported_at)
    return _save_document(document)


def export_participation_docx(ctx: TenderExportContext) -> bytes:
    analysis = ctx.participation
    document = _setup_document(
        "Report partecipabilità",
        f"Gara CIG {ctx.tender.cig} · {analysis.forma_label}",
    )

    copertura = analysis.copertura
    document.add_paragraph(
        f"Copertura requisiti: {copertura.percentuale:.1f}% "
        f"({copertura.soddisfatti}/{copertura.totale} soddisfatti)"
    )

    if ctx.participation_suggestion:
        suggestion = ctx.participation_suggestion
        document.add_paragraph(
            f"Suggerimento sistema: {suggestion.get('forma_label', '')} "
            f"({suggestion.get('confidenza', '')}) — {suggestion.get('motivazione', '')}"
        )

    if analysis.criticita:
        document.add_heading("Criticità", level=2)
        for item in analysis.criticita:
            document.add_paragraph(
                f"[{item.get('severita', '').upper()}] {item.get('descrizione', '')}: "
                f"{item.get('motivazione', '')}",
                style="List Bullet",
            )

    document.add_heading("Dettaglio requisiti", level=2)
    table = document.add_table(rows=1 + len(analysis.requisiti), cols=6)
    table.style = "Table Grid"
    headers = ["Requisito", "Esito", "Azienda", "Posseduto", "Richiesto", "Motivazione"]
    for col_index, header in enumerate(headers):
        table.rows[0].cells[col_index].text = header
        for run in table.rows[0].cells[col_index].paragraphs[0].runs:
            run.bold = True

    for row_index, req in enumerate(analysis.requisiti, start=1):
        row = table.rows[row_index].cells
        row[0].text = req.descrizione[:300]
        row[1].text = req.esito_label
        row[2].text = req.company_name or "—"
        row[3].text = req.valore_posseduto
        row[4].text = req.valore_richiesto
        row[5].text = req.motivazione

    _add_meta_footer(document, ctx.exported_at)
    return _save_document(document)


def export_relation_docx(ctx: TenderExportContext) -> bytes:
    relation = ctx.relation
    company = relation.get("company_name") or "Non specificata"
    document = _setup_document(
        "Relazione tecnica",
        f"Gara CIG {ctx.tender.cig} · Azienda: {company}",
    )

    sections = sorted(relation.get("sections") or [], key=lambda item: item.get("order", 0))
    if not sections:
        document.add_paragraph("Nessuna sezione compilata nella relazione tecnica.")
    else:
        for section in sections:
            document.add_heading(section.get("title", "Sezione"), level=2)
            category = section.get("category", "")
            if category:
                document.add_paragraph(f"Categoria: {category}")
            pages = section.get("suggested_pages")
            if pages:
                document.add_paragraph(f"Pagine suggerite: {pages}")
            content = _markdown_to_plain(section.get("content", ""))
            for paragraph in content.split("\n\n"):
                if paragraph.strip():
                    document.add_paragraph(paragraph.strip())

    _add_meta_footer(document, ctx.exported_at)
    return _save_document(document)


def export_economic_docx(ctx: TenderExportContext) -> bytes:
    economic = ctx.economic_relation
    company = economic.get("company_name") or "Non specificata"
    outline = economic.get("outline") or {}
    document = _setup_document(
        "Offerta economica",
        f"Gara CIG {ctx.tender.cig} · Azienda: {company}",
    )
    document.add_paragraph(
        f"Modello: {outline.get('pricing_model', '—')} · "
        f"Importo base: {outline.get('importo_base', '—')}"
    )

    items = sorted(economic.get("line_items") or [], key=lambda item: item.get("order", 0))
    if not items:
        document.add_paragraph("Nessuna voce economica compilata.")
    else:
        table = document.add_table(rows=len(items) + 1, cols=6)
        table.style = "Table Grid"
        headers = ("Voce", "U.M.", "Q.tà", "Prezzo unit.", "Importo", "Ribasso %")
        for col_index, header in enumerate(headers):
            table.rows[0].cells[col_index].text = header
        for row_index, item in enumerate(items, start=1):
            values = (
                item.get("voce", ""),
                item.get("unita_misura", ""),
                str(item.get("quantita", "")),
                str(item.get("prezzo_unitario", "")),
                str(item.get("importo", "")),
                str(item.get("ribasso_percentuale", "")),
            )
            for col_index, value in enumerate(values):
                table.rows[row_index].cells[col_index].text = value

    totals = outline.get("totals") or {}
    if totals.get("totale"):
        document.add_paragraph(f"Totale offerta: € {totals.get('totale')}")

    _add_meta_footer(document, ctx.exported_at)
    return _save_document(document)


def _markdown_to_plain(text: str) -> str:
    cleaned = re.sub(r"^#+\s*", "", text, flags=re.MULTILINE)
    cleaned = re.sub(r"\*\*(.*?)\*\*", r"\1", cleaned)
    cleaned = re.sub(r"\*(.*?)\*", r"\1", cleaned)
    return cleaned


def build_docx_export(item: str, ctx: TenderExportContext) -> bytes:
    builders = {
        SCHEDA_GARA: export_scheda_docx,
        MATRICE_REQUISITI: export_matrix_docx,
        REPORT_PARTECIPABILITA: export_participation_docx,
        RELAZIONE_TECNICA: export_relation_docx,
        OFFERTA_ECONOMICA: export_economic_docx,
    }
    return builders[item](ctx)
